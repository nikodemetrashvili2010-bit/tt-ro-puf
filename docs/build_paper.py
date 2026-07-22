#!/usr/bin/env python3
"""Build the repository paper from docs/paper_draft.md.

The Markdown file is the source of truth. This script writes matching DOCX and
PDF editions with inline figures, document metadata, and page numbers. It uses
only python-docx, Pillow, and ReportLab.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import os
import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as PDFImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "paper_draft.md"
OUTPUT_DIR = ROOT / "docs" / "paper"
DOCX_PATH = OUTPUT_DIR / "SILICON_paper.docx"
PDF_PATH = OUTPUT_DIR / "SILICON_paper.pdf"
BUILD_SCHEMA = "paper-builder-v5"
DOCUMENT_DATE = datetime(2026, 7, 22, tzinfo=timezone.utc)
ZIP_DATE = (2026, 7, 22, 0, 0, 0)

INLINE_TOKEN = re.compile(
    r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`|\[[^\]]+\]\([^)]+\))"
)
IMAGE_LINE = re.compile(r"^!\[(?P<alt>.*)\]\((?P<path>[^)]+)\)$")


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    path: str = ""


def parse_markdown(path: Path) -> list[Block]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[Block] = []
    paragraph: list[str] = []
    list_lines: list[str] = []
    list_kind = ""
    list_number = 0

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(Block("paragraph", " ".join(x.strip() for x in paragraph)))
            paragraph.clear()

    def flush_list() -> None:
        nonlocal list_kind, list_number
        if list_lines:
            blocks.append(
                Block(list_kind, " ".join(x.strip() for x in list_lines), list_number)
            )
            list_lines.clear()
            list_kind = ""
            list_number = 0

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            flush_list()
            continue
        if stripped == "---":
            flush_paragraph()
            flush_list()
            blocks.append(Block("rule"))
            continue
        image_match = IMAGE_LINE.match(stripped)
        if image_match:
            flush_paragraph()
            flush_list()
            blocks.append(
                Block("image", image_match.group("alt"), path=image_match.group("path"))
            )
            continue
        if stripped.startswith("#"):
            heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if heading:
                flush_paragraph()
                flush_list()
                blocks.append(Block("heading", heading.group(2), len(heading.group(1))))
                continue
        # Markdown list ordinals are short. A four-digit year at the start of
        # a wrapped bibliography line (for example, "2021.") is prose.
        ordered = re.match(r"^(\d{1,2})\.\s+(.*)$", stripped)
        if ordered:
            flush_paragraph()
            flush_list()
            list_kind = "number"
            list_number = int(ordered.group(1))
            list_lines.append(ordered.group(2))
            continue
        if line.startswith("- "):
            flush_paragraph()
            flush_list()
            list_kind = "bullet"
            list_lines.append(line[2:])
            continue
        if list_lines and (line.startswith("  ") or line.startswith("\t")):
            list_lines.append(stripped)
            continue
        if list_lines:
            flush_list()
        paragraph.append(stripped)

    flush_paragraph()
    flush_list()
    return blocks


def plain_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def resolve_asset(path: str) -> Path:
    candidate = (SOURCE.parent / path).resolve()
    try:
        candidate.relative_to(ROOT)
    except ValueError as exc:
        raise ValueError(f"Figure path leaves the repository: {path}") from exc
    return candidate


def paper_build_id(blocks: list[Block]) -> str:
    digest = hashlib.sha256()
    digest.update(BUILD_SCHEMA.encode("utf-8"))
    digest.update(b"\0")
    digest.update(SOURCE.read_bytes())
    for block in blocks:
        if block.kind == "image":
            image_path = resolve_asset(block.path)
            digest.update(block.path.encode("utf-8"))
            digest.update(image_path.read_bytes())
    return digest.hexdigest()


def add_docx_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        run = paragraph.add_run()
        if token.startswith("**"):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith("*"):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith("`"):
            run.text = token[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            assert link
            run.text = f"{link.group(1)} ({link.group(2)})"
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def set_cell_font(style, name: str, size: float, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, value, end):
        run._r.append(element)


def add_bottom_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), "555555")
    borders.append(bottom)
    p_pr.append(borders)


def fit_image(path: Path, max_width: float, max_height: float) -> tuple[float, float]:
    with PILImage.open(path) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    return width * scale, height * scale


def build_docx(blocks: list[Block], title: str, author: str, build_id: str) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = document.styles
    normal = styles["Normal"]
    set_cell_font(normal, "Times New Roman", 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.06

    title_style = styles["Title"]
    set_cell_font(title_style, "Arial", 18, True)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_style.paragraph_format.space_after = Pt(10)

    for name, size, before, after in (
        ("Heading 1", 13.5, 10, 5),
        ("Heading 2", 11.5, 8, 4),
        ("Heading 3", 10.5, 6, 3),
    ):
        style = styles[name]
        set_cell_font(style, "Arial", size, True)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    set_cell_font(caption, "Times New Roman", 9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(55, 55, 55)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    if "Bibliography" not in styles:
        bibliography = styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bibliography = styles["Bibliography"]
    set_cell_font(bibliography, "Times New Roman", 9)
    bibliography.paragraph_format.left_indent = Inches(0.24)
    bibliography.paragraph_format.first_line_indent = Inches(-0.24)
    bibliography.paragraph_format.space_after = Pt(4)
    bibliography.paragraph_format.line_spacing = 1.0
    bibliography.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    core = document.core_properties
    core.title = title
    core.author = author
    core.subject = "Pre-silicon study of layout-induced bias in a sky130 ring-oscillator PUF"
    core.keywords = (
        "RO-PUF, sky130, TinyTapeout, parasitics, pre-silicon, "
        f"build-sha256={build_id}"
    )
    core.comments = "Built from docs/paper_draft.md and the four embedded figures"
    core.created = DOCUMENT_DATE
    core.modified = DOCUMENT_DATE

    header = section.header.paragraphs[0]
    header.text = "SILICON - pre-silicon RO-PUF study"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
    add_page_field(section.footer.paragraphs[0])

    in_references = False
    for index, block in enumerate(blocks):
        if block.kind == "heading" and block.level == 1:
            paragraph = document.add_paragraph(style="Title")
            add_docx_inline(paragraph, block.text)
            continue
        if block.kind == "rule":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            add_bottom_rule(paragraph)
            continue
        if block.kind == "heading":
            level = min(max(block.level - 1, 1), 3)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_docx_inline(paragraph, block.text)
            in_references = plain_markdown(block.text).strip().lower() == "references"
            continue
        if block.kind == "image":
            image_path = resolve_asset(block.path)
            if not image_path.is_file():
                raise FileNotFoundError(f"Figure not found: {image_path}")
            width_px, height_px = fit_image(image_path, 6.25 * 100, 4.15 * 100)
            picture_paragraph = document.add_paragraph()
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = picture_paragraph.add_run()
            shape = run.add_picture(
                str(image_path), width=Inches(width_px / 100), height=Inches(height_px / 100)
            )
            shape._inline.docPr.set("descr", plain_markdown(block.text))
            shape._inline.docPr.set("title", plain_markdown(block.text).split(":", 1)[0])
            caption_paragraph = document.add_paragraph(style="Caption")
            add_docx_inline(caption_paragraph, block.text)
            continue
        if block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.34)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(3)
            add_docx_inline(paragraph, block.text)
            continue
        if block.kind == "number":
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.34)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(3)
            add_docx_inline(paragraph, block.text)
            continue
        if block.kind == "paragraph":
            style = "Bibliography" if in_references and re.match(r"^\[\d+\]", block.text) else None
            paragraph = document.add_paragraph(style=style)
            if index in (1, 2):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(7)
            add_docx_inline(paragraph, block.text)

    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)


def register_pdf_fonts() -> tuple[str, str, str, str]:
    candidates = [
        (
            Path("C:/Windows/Fonts/times.ttf"),
            Path("C:/Windows/Fonts/timesbd.ttf"),
            Path("C:/Windows/Fonts/timesi.ttf"),
            Path("C:/Windows/Fonts/timesbi.ttf"),
        ),
        (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            Path("C:/Windows/Fonts/ariali.ttf"),
            Path("C:/Windows/Fonts/arialbi.ttf"),
        ),
    ]
    for regular, bold, italic, bold_italic in candidates:
        if all(path.is_file() for path in (regular, bold, italic, bold_italic)):
            pdfmetrics.registerFont(TTFont("Paper", regular))
            pdfmetrics.registerFont(TTFont("Paper-Bold", bold))
            pdfmetrics.registerFont(TTFont("Paper-Italic", italic))
            pdfmetrics.registerFont(TTFont("Paper-BoldItalic", bold_italic))
            pdfmetrics.registerFontFamily(
                "Paper",
                normal="Paper",
                bold="Paper-Bold",
                italic="Paper-Italic",
                boldItalic="Paper-BoldItalic",
            )
            return "Paper", "Paper-Bold", "Paper-Italic", "Paper-BoldItalic"
    return "Times-Roman", "Times-Bold", "Times-Italic", "Times-BoldItalic"


def pdf_inline(text: str) -> str:
    pieces: list[str] = []
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        pieces.append(html.escape(text[position : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            pieces.append(f"<b>{html.escape(token[2:-2])}</b>")
        elif token.startswith("*"):
            pieces.append(f"<i>{html.escape(token[1:-1])}</i>")
        elif token.startswith("`"):
            pieces.append(html.escape(token[1:-1]))
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            assert link
            pieces.append(
                f'<link href="{html.escape(link.group(2), quote=True)}">'
                f"{html.escape(link.group(1))}</link>"
            )
        position = match.end()
    pieces.append(html.escape(text[position:]))
    return "".join(pieces)


def build_pdf(blocks: list[Block], title: str, author: str, build_id: str) -> int:
    regular, bold, italic, _ = register_pdf_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PaperBody",
        parent=styles["BodyText"],
        fontName=regular,
        fontSize=10.5,
        leading=12.1,
        alignment=TA_JUSTIFY,
        spaceAfter=5,
        allowWidows=0,
        allowOrphans=0,
    )
    title_style = ParagraphStyle(
        "PaperTitle",
        parent=body,
        fontName=bold,
        fontSize=18,
        leading=22,
        alignment=TA_LEFT,
        spaceAfter=10,
    )
    meta = ParagraphStyle(
        "PaperMeta", parent=body, fontSize=11.5, leading=14, alignment=TA_LEFT, spaceAfter=7
    )
    headings = {
        2: ParagraphStyle(
            "PaperH1", parent=body, fontName=bold, fontSize=13.5, leading=16,
            spaceBefore=11, spaceAfter=5, keepWithNext=True,
        ),
        3: ParagraphStyle(
            "PaperH2", parent=body, fontName=bold, fontSize=11.5, leading=14,
            spaceBefore=8, spaceAfter=4, keepWithNext=True,
        ),
        4: ParagraphStyle(
            "PaperH3", parent=body, fontName=bold, fontSize=10.5, leading=13,
            spaceBefore=7, spaceAfter=3, keepWithNext=True,
        ),
    }
    bullet = ParagraphStyle(
        "PaperBullet", parent=body, leftIndent=18, firstLineIndent=-10, bulletIndent=3,
        spaceAfter=3,
    )
    caption = ParagraphStyle(
        "PaperCaption", parent=body, fontName=italic, fontSize=9, leading=10.5,
        alignment=TA_CENTER, textColor="#333333", spaceBefore=3, spaceAfter=8,
    )
    bibliography = ParagraphStyle(
        "PaperBibliography", parent=body, fontSize=9, leading=10.6,
        leftIndent=17, firstLineIndent=-17, spaceAfter=4, alignment=TA_LEFT,
    )

    page_count = 0

    def draw_page(canvas, doc) -> None:
        nonlocal page_count
        page_count = max(page_count, doc.page)
        canvas.saveState()
        canvas.setTitle(title)
        canvas.setAuthor(author)
        canvas.setSubject("Pre-silicon study of layout-induced bias in a sky130 RO-PUF")
        canvas.setKeywords(
            "RO-PUF, sky130, TinyTapeout, pre-silicon, "
            f"build-sha256={build_id}"
        )
        canvas.setFont(regular, 8)
        canvas.setFillColorRGB(0.38, 0.38, 0.38)
        if doc.page > 1:
            canvas.drawRightString(7.68 * inch, 10.63 * inch, "SILICON - pre-silicon RO-PUF study")
        canvas.drawCentredString(4.25 * inch, 0.34 * inch, str(doc.page))
        canvas.restoreState()

    story = []
    in_references = False
    for index, block in enumerate(blocks):
        if block.kind == "heading" and block.level == 1:
            story.append(Paragraph(pdf_inline(block.text), title_style))
            continue
        if block.kind == "rule":
            story.append(Spacer(1, 5))
            continue
        if block.kind == "heading":
            style = headings.get(min(block.level, 4), headings[4])
            story.append(Paragraph(pdf_inline(block.text), style))
            in_references = plain_markdown(block.text).strip().lower() == "references"
            continue
        if block.kind == "image":
            image_path = resolve_asset(block.path)
            if not image_path.is_file():
                raise FileNotFoundError(f"Figure not found: {image_path}")
            width, height = fit_image(image_path, 6.25 * 100, 4.0 * 100)
            figure = PDFImage(str(image_path), width=width / 100 * inch, height=height / 100 * inch)
            figure.hAlign = "CENTER"
            story.append(KeepTogether([figure, Paragraph(pdf_inline(block.text), caption)]))
            continue
        if block.kind == "bullet":
            story.append(Paragraph(pdf_inline(block.text), bullet, bulletText="•"))
            continue
        if block.kind == "number":
            story.append(
                Paragraph(pdf_inline(block.text), bullet, bulletText=f"{block.level}.")
            )
            continue
        if block.kind == "paragraph":
            style = body
            if index in (1, 2):
                style = meta
            elif in_references and re.match(r"^\[\d+\]", block.text):
                style = bibliography
            story.append(Paragraph(pdf_inline(block.text), style))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        invariant=1,
        rightMargin=0.82 * inch,
        leftMargin=0.82 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.62 * inch,
        title=title,
        author=author,
        subject="Pre-silicon study of layout-induced bias in a sky130 RO-PUF",
    )
    source_date_epoch = str(int(DOCUMENT_DATE.timestamp()))
    previous_source_date_epoch = os.environ.get("SOURCE_DATE_EPOCH")
    os.environ["SOURCE_DATE_EPOCH"] = source_date_epoch
    try:
        document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    finally:
        if previous_source_date_epoch is None:
            os.environ.pop("SOURCE_DATE_EPOCH", None)
        else:
            os.environ["SOURCE_DATE_EPOCH"] = previous_source_date_epoch
    return page_count


def patch_docx_app_properties(path: Path, pages: int, words: int) -> None:
    with tempfile.TemporaryDirectory(prefix="paper_docx_") as temp_dir:
        temp = Path(temp_dir)
        with zipfile.ZipFile(path, "r") as source:
            source.extractall(temp)
        app_path = temp / "docProps" / "app.xml"
        if app_path.is_file():
            text = app_path.read_text(encoding="utf-8")
            text = re.sub(r"<Pages>\d+</Pages>", f"<Pages>{pages}</Pages>", text)
            text = re.sub(r"<Words>\d+</Words>", f"<Words>{words}</Words>", text)
            app_path.write_text(text, encoding="utf-8")

        # The default Word template carries editing-session identifiers and an
        # empty bibliography custom-XML part. Neither belongs in a published
        # source-built artifact, so remove them before packaging.
        custom_xml = temp / "customXml"
        if custom_xml.is_dir():
            shutil.rmtree(custom_xml)
        content_types = temp / "[Content_Types].xml"
        text = content_types.read_text(encoding="utf-8")
        text = re.sub(
            r'<Override\b[^>]*PartName="/customXml/[^\"]+"[^>]*/>', "", text
        )
        content_types.write_text(text, encoding="utf-8")
        relationships = temp / "word" / "_rels" / "document.xml.rels"
        text = relationships.read_text(encoding="utf-8")
        text = re.sub(
            r'<Relationship\b[^>]*Type="[^\"]*/customXml"[^>]*/>', "", text
        )
        relationships.write_text(text, encoding="utf-8")

        for xml_path in temp.rglob("*.xml"):
            text = xml_path.read_text(encoding="utf-8")
            text = re.sub(r'\s+w:rsid\w*="[^\"]*"', "", text)
            text = re.sub(r"<w:rsids\b[^>]*>.*?</w:rsids>", "", text, flags=re.DOTALL)
            text = re.sub(r"<w:rsid(?:Root)?\b[^>]*/>", "", text)
            text = re.sub(r"<w14:docId\b[^>]*/>", "", text)
            xml_path.write_text(text, encoding="utf-8")

        rebuilt = path.with_suffix(".rebuilt.docx")
        with zipfile.ZipFile(rebuilt, "w", zipfile.ZIP_DEFLATED) as target:
            for item in sorted(temp.rglob("*")):
                if item.is_file():
                    info = zipfile.ZipInfo(item.relative_to(temp).as_posix(), ZIP_DATE)
                    info.compress_type = zipfile.ZIP_DEFLATED
                    info.external_attr = 0o600 << 16
                    target.writestr(info, item.read_bytes())
        rebuilt.replace(path)


def main() -> None:
    global SOURCE, OUTPUT_DIR, DOCX_PATH, PDF_PATH
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()

    SOURCE = args.source.resolve()
    OUTPUT_DIR = args.output_dir.resolve()
    DOCX_PATH = OUTPUT_DIR / "SILICON_paper.docx"
    PDF_PATH = OUTPUT_DIR / "SILICON_paper.pdf"

    blocks = parse_markdown(SOURCE)
    if not blocks or blocks[0].kind != "heading" or blocks[0].level != 1:
        raise SystemExit("paper_draft.md must begin with one level-1 title")
    title = plain_markdown(blocks[0].text)
    author = "Nikoloz Demetrashvili"
    build_id = paper_build_id(blocks)
    build_docx(blocks, title, author, build_id)
    pages = build_pdf(blocks, title, author, build_id)
    words = len(re.findall(r"\b[\w'-]+\b", plain_markdown(SOURCE.read_text(encoding="utf-8"))))
    patch_docx_app_properties(DOCX_PATH, pages, words)
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH} ({pages} pages, {words} words)")
    print(f"Build SHA-256: {build_id}")


if __name__ == "__main__":
    main()
